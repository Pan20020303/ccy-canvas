from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


DESKTOP = Path(r"C:\Users\19536\Desktop")
OUTPUT_DOCX = DESKTOP / "\u7535\u5546\u5e26\u8d27\u89c6\u9891\u9879\u76ee\u6c47\u62a5.docx"

TEXT = {
    "title": "\u7535\u5546\u5e26\u8d27\u89c6\u9891\u9879\u76ee\u6c47\u62a5",
    "subtitle": "\u57fa\u4e8e\u9879\u76ee\u754c\u9762\u622a\u56fe\u6574\u7406\u7684\u767d\u8bdd\u7248\u8bf4\u660e",
    "intro": "\u8fd9\u4efd\u6750\u6599\u4e3b\u8981\u56de\u7b54\u4e09\u4e2a\u95ee\u9898\uff1a\u9879\u76ee\u662f\u505a\u4ec0\u4e48\u7684\u3001\u73b0\u5728\u505a\u5230\u54ea\u4e00\u6b65\u3001\u80fd\u5e26\u6765\u4ec0\u4e48\u4ef7\u503c\u3002",
    "home_caption": "\u9879\u76ee\u9996\u9875\uff1a\u5df2\u7ecf\u5177\u5907\u65b0\u5efa\u89c6\u9891\u3001\u5546\u54c1\u5e93\u3001\u7206\u6b3e\u590d\u523b\u548c\u9879\u76ee\u7ba1\u7406\u5165\u53e3",
    "s1": "\u4e00\u3001\u9879\u76ee\u662f\u505a\u4ec0\u4e48\u7684",
    "p1": "\u8fd9\u662f\u4e00\u4e2a\u7528 AI \u5e2e\u5546\u5bb6\u5feb\u901f\u751f\u6210\u5e26\u8d27\u77ed\u89c6\u9891\u7684\u5de5\u5177\u3002\u7528\u6237\u5f55\u5165\u5546\u54c1\u4fe1\u606f\u540e\uff0c\u7cfb\u7edf\u53ef\u4ee5\u7ee7\u7eed\u5b8c\u6210\u811a\u672c\u751f\u6210\u3001\u7d20\u6750\u751f\u6210\u3001\u89c6\u9891\u5408\u6210\u548c\u5bfc\u51fa\u3002",
    "p2": "\u901a\u4fd7\u5730\u8bf4\uff0c\u8fd9\u4e2a\u9879\u76ee\u5c31\u662f\u628a\u539f\u6765\u5f88\u4f9d\u8d56\u4eba\u5de5\u7ecf\u9a8c\u7684\u5e26\u8d27\u89c6\u9891\u5236\u4f5c\u6d41\u7a0b\uff0c\u5c3d\u91cf\u53d8\u6210\u4e00\u6761\u6807\u51c6\u5316\u6d41\u6c34\u7ebf\u3002",
    "s2": "\u4e8c\u3001\u76ee\u524d\u5df2\u7ecf\u505a\u6210\u7684\u6838\u5fc3\u80fd\u529b",
    "b1": "1. \u9996\u9875\u5df2\u7ecf\u5f62\u6210\u5b8c\u6574\u5165\u53e3\uff0c\u7528\u6237\u53ef\u4ee5\u65b0\u5efa\u89c6\u9891\u9879\u76ee\uff0c\u4e5f\u53ef\u4ee5\u8fdb\u5165\u5546\u54c1\u5e93\u548c\u7206\u6b3e\u590d\u523b\u6a21\u5757\u3002",
    "b2": "2. \u811a\u672c\u9875\u652f\u6301 AI \u4e00\u6b21\u751f\u6210\u591a\u4e2a\u65b9\u6848\uff0c\u65b9\u4fbf\u7528\u6237\u5148\u9009\u65b9\u5411\uff0c\u518d\u505a\u5fae\u8c03\u3002",
    "b3": "3. \u7d20\u6750\u9875\u4f1a\u6309\u7167\u811a\u672c\u62c6\u5206\u5206\u955c\uff0c\u9010\u6761\u751f\u6210\u6240\u9700\u753b\u9762\uff0c\u51cf\u5c11\u4eba\u5de5\u627e\u7d20\u6750\u7684\u65f6\u95f4\u3002",
    "b4": "4. \u5408\u6210\u9875\u652f\u6301\u914d\u97f3\u3001\u80cc\u666f\u97f3\u4e50\u3001\u5b57\u5e55\u3001\u753b\u9762\u6bd4\u4f8b\u548c\u5206\u8fa8\u7387\u8bbe\u7f6e\uff0c\u8bf4\u660e\u7cfb\u7edf\u5df2\u7ecf\u8986\u76d6\u5230\u51fa\u6210\u7247\u8fd9\u4e00\u6b65\u3002",
    "b5": "5. \u5bfc\u51fa\u9875\u652f\u6301\u591a\u5e73\u53f0\u7248\u672c\u8f93\u51fa\uff0c\u8fd8\u5e26\u6709 A/B \u6d4b\u8bd5\u7248\u672c\u80fd\u529b\uff0c\u66f4\u63a5\u8fd1\u771f\u5b9e\u6295\u653e\u573a\u666f\u3002",
    "script_caption": "\u811a\u672c\u65b9\u6848\u9875\uff1a\u540c\u4e00\u4e2a\u5546\u54c1\u53ef\u4ee5\u5feb\u901f\u5c1d\u8bd5\u591a\u4e2a\u5e26\u8d27\u8868\u8fbe\u65b9\u5411",
    "asset_caption": "\u7d20\u6750\u751f\u6210\u9875\uff1a\u6309\u5206\u955c\u9010\u6761\u8865\u9f50\u753b\u9762\u7d20\u6750\uff0c\u51cf\u5c11\u91cd\u590d\u4eba\u5de5\u64cd\u4f5c",
    "compose_caption": "\u89c6\u9891\u5408\u6210\u9875\uff1a\u53ef\u4ee5\u76f4\u63a5\u914d\u7f6e\u914d\u97f3\u3001\u5b57\u5e55\u3001\u6bd4\u4f8b\u548c\u6e05\u6670\u5ea6",
    "export_caption": "\u5bfc\u51fa\u9875\uff1a\u652f\u6301\u591a\u5e73\u53f0\u9002\u914d\u548c\u7ee7\u7eed\u751f\u6210 A/B \u6d4b\u8bd5\u7248\u672c",
    "s3": "\u4e09\u3001\u8fd9\u4e2a\u9879\u76ee\u80fd\u5e26\u6765\u4ec0\u4e48\u4ef7\u503c",
    "v1": "\u7b2c\u4e00\uff0c\u63d0\u6548\u3002\u539f\u6765\u5206\u6563\u5728\u6587\u6848\u3001\u627e\u56fe\u3001\u526a\u8f91\u91cc\u7684\u52a8\u4f5c\uff0c\u88ab\u4e32\u6210\u4e00\u6761\u6d41\u7a0b\uff0c\u80fd\u660e\u663e\u7f29\u77ed\u51fa\u7247\u65f6\u95f4\u3002",
    "v2": "\u7b2c\u4e8c\uff0c\u964d\u95e8\u69db\u3002\u4e0d\u64c5\u957f\u811a\u672c\u548c\u526a\u8f91\u7684\u4eba\uff0c\u4e5f\u80fd\u901a\u8fc7\u7cfb\u7edf\u5feb\u901f\u505a\u51fa\u53ef\u7528\u7248\u672c\u3002",
    "v3": "\u7b2c\u4e09\uff0c\u4fbf\u4e8e\u590d\u5236\u3002\u5546\u54c1\u5e93\u3001\u7206\u6b3e\u590d\u523b\u548c A/B \u6d4b\u8bd5\uff0c\u8bf4\u660e\u597d\u7684\u5185\u5bb9\u53ef\u4ee5\u4e0d\u65ad\u590d\u7528\u3002",
    "v4": "\u7b2c\u56db\uff0c\u66f4\u8d34\u8fd1\u4e1a\u52a1\u3002\u591a\u5e73\u53f0\u5bfc\u51fa\u8bf4\u660e\u9879\u76ee\u5df2\u7ecf\u5728\u8003\u8651\u771f\u5b9e\u8fd0\u8425\u548c\u6295\u653e\uff0c\u800c\u4e0d\u662f\u53ea\u505a\u6f14\u793a\u3002",
    "s4": "\u56db\u3001\u5f53\u524d\u8fdb\u5c55\u5224\u65ad",
    "p3": "\u4ece\u622a\u56fe\u770b\uff0c\u8fd9\u4e2a\u9879\u76ee\u5df2\u7ecf\u628a\u5efa\u9879\u76ee\u3001\u5199\u811a\u672c\u3001\u4ea7\u7d20\u6750\u3001\u5408\u6210\u89c6\u9891\u3001\u5bfc\u51fa\u7248\u672c\u8fd9\u6761\u4e3b\u94fe\u8def\u8dd1\u901a\u4e86\u3002\u5b83\u5df2\u7ecf\u4e0d\u662f\u4e00\u4e2a\u5355\u70b9\u529f\u80fd\uff0c\u800c\u662f\u4e00\u5957\u8f83\u5b8c\u6574\u7684\u7535\u5546\u5e26\u8d27\u89c6\u9891\u751f\u4ea7\u5de5\u4f5c\u53f0\u3002",
    "p4": "\u5982\u679c\u540e\u7eed\u7ee7\u7eed\u63d0\u5347\u751f\u6210\u8d28\u91cf\u3001\u8865\u5145\u6548\u679c\u53cd\u9988\u548c\u534f\u4f5c\u80fd\u529b\uff0c\u8fd9\u5957\u7cfb\u7edf\u6709\u673a\u4f1a\u4ece\u5236\u4f5c\u5de5\u5177\u5347\u7ea7\u6210\u7535\u5546\u5185\u5bb9\u751f\u4ea7\u5e73\u53f0\u3002",
    "s5": "\u4e94\u3001\u5efa\u8bae\u4e0b\u4e00\u6b65\u91cd\u70b9",
    "n1": "1. \u7ee7\u7eed\u4f18\u5316\u811a\u672c\u81ea\u7136\u5ea6\u3001\u7d20\u6750\u4e00\u81f4\u6027\u548c\u6210\u7247\u89c2\u611f\u3002",
    "n2": "2. \u589e\u52a0\u6548\u679c\u53cd\u9988\u80fd\u529b\uff0c\u5e2e\u52a9\u56e2\u961f\u5224\u65ad\u54ea\u79cd\u811a\u672c\u548c\u7248\u672c\u8f6c\u5316\u66f4\u597d\u3002",
    "n3": "3. \u5b8c\u5584\u5546\u54c1\u5e93\u548c\u590d\u523b\u80fd\u529b\uff0c\u8ba9\u9ad8\u8868\u73b0\u5185\u5bb9\u53ef\u4ee5\u5feb\u901f\u8fc1\u79fb\u5230\u76f8\u4f3c\u5546\u54c1\u3002",
    "n4": "4. \u52a0\u5f3a\u56e2\u961f\u534f\u4f5c\u548c\u5ba1\u6838\u6d41\u7a0b\uff0c\u652f\u6301\u591a\u4eba\u5171\u540c\u4f7f\u7528\u3002",
    "summary_h": "\u603b\u7ed3",
    "summary_p": "\u6574\u4f53\u6765\u770b\uff0c\u8fd9\u4e2a\u9879\u76ee\u65b9\u5411\u6e05\u6670\uff0c\u4ef7\u503c\u660e\u786e\uff0c\u6838\u5fc3\u6d41\u7a0b\u4e5f\u5df2\u7ecf\u5177\u5907\u96cf\u5f62\u3002\u5b83\u6700\u5927\u7684\u610f\u4e49\uff0c\u662f\u628a\u7535\u5546\u5e26\u8d27\u89c6\u9891\u4ece\u96f6\u6563\u5236\u4f5c\uff0c\u63a8\u8fdb\u5230\u6807\u51c6\u5316\u3001\u81ea\u52a8\u5316\u751f\u4ea7\u3002",
}

IMAGES = {
    "home": DESKTOP / "e6dc3534-6ebf-4fe4-8c81-768ac6bdf2de.png",
    "script": DESKTOP / "5f45bb5f-7753-46f4-8839-ccc851ddcc52.png",
    "asset": DESKTOP / "5a554ab7-d7c4-43e0-a418-c2c07d6fc366.png",
    "compose": DESKTOP / "daf0b4b8-0352-4194-9c35-b8f15d683ccb.png",
    "export": DESKTOP / "f1e5a7d4-d160-4236-b716-63f9e1ce15a2.png",
}


def set_run_font(run, size=11, bold=False, color=None):
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for key in ("w:ascii", "w:hAnsi", "w:eastAsia"):
        r_fonts.set(qn(key), "Microsoft YaHei")


def add_paragraph(doc, text, size=11, bold=False, color=None, before=0, after=8, line=1.25, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    size = 18 if level == 1 else 14
    before = 14 if level == 1 else 8
    after = 6 if level == 1 else 4
    add_paragraph(
        doc,
        text,
        size=size,
        bold=True,
        color=RGBColor(46, 116, 181),
        before=before,
        after=after,
        line=1.15,
    )


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run)


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    set_run_font(run)


def add_image(doc, path, caption):
    doc.add_picture(str(path), width=Inches(5.75))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(
        doc,
        caption,
        size=9.5,
        color=RGBColor(95, 95, 95),
        before=2,
        after=10,
        line=1.0,
        align=WD_ALIGN_PARAGRAPH.CENTER,
    )


def configure(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def build_cover(doc):
    add_paragraph(doc, TEXT["title"], size=24, bold=True, before=46, after=6, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, TEXT["subtitle"], size=13, color=RGBColor(85, 85, 85), before=0, after=20, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_paragraph(doc, TEXT["intro"], size=12, before=0, after=14, line=1.35, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_image(doc, IMAGES["home"], TEXT["home_caption"])
    add_paragraph(doc, "2026-05-29", size=10.5, color=RGBColor(110, 110, 110), before=16, after=0, line=1.0, align=WD_ALIGN_PARAGRAPH.CENTER)


def build_body(doc):
    add_heading(doc, TEXT["s1"])
    add_paragraph(doc, TEXT["p1"], line=1.35)
    add_paragraph(doc, TEXT["p2"], line=1.35)

    add_heading(doc, TEXT["s2"])
    for key in ("b1", "b2", "b3", "b4", "b5"):
        add_bullet(doc, TEXT[key])
    add_image(doc, IMAGES["script"], TEXT["script_caption"])
    add_image(doc, IMAGES["asset"], TEXT["asset_caption"])
    add_image(doc, IMAGES["compose"], TEXT["compose_caption"])
    add_image(doc, IMAGES["export"], TEXT["export_caption"])

    add_heading(doc, TEXT["s3"])
    for key in ("v1", "v2", "v3", "v4"):
        add_bullet(doc, TEXT[key])

    add_heading(doc, TEXT["s4"])
    add_paragraph(doc, TEXT["p3"], line=1.35)
    add_paragraph(doc, TEXT["p4"], line=1.35)

    add_heading(doc, TEXT["s5"])
    for key in ("n1", "n2", "n3", "n4"):
        add_number(doc, TEXT[key])

    add_heading(doc, TEXT["summary_h"])
    add_paragraph(doc, TEXT["summary_p"], line=1.35)


def main():
    for path in IMAGES.values():
        if not path.exists():
            raise FileNotFoundError(f"Missing image: {path}")
    doc = Document()
    configure(doc)
    build_cover(doc)
    doc.add_section(WD_SECTION_START.NEW_PAGE)
    build_body(doc)
    doc.save(OUTPUT_DOCX)
    print(OUTPUT_DOCX)


if __name__ == "__main__":
    main()
